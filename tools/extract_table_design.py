from pathlib import Path
import re

import docx


SOURCE = Path("01-详细设计文档1024.docx")
OUT_DIR = Path("表结构设计")

MODULE_ORDER = [
    "OEE与KPI分析",
    "设备资产管理",
    "预防性维护-点巡检",
    "预防性维护-设备保养",
    "故障维修与异常工单",
    "工装管理（详细设计来源）",
    "备件管理",
]


def iter_blocks(document):
    para_index = 0
    table_index = 0
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield ("p", para_index, document.paragraphs[para_index].text.strip())
            para_index += 1
        elif tag == "tbl":
            yield ("t", table_index, "")
            table_index += 1


def module_from_heading(text):
    for pattern, module in [
        (r"^3\.1\b", "OEE与KPI分析"),
        (r"^3\.2\b", "设备资产管理"),
        (r"^3\.3\b", "预防性维护-点巡检"),
        (r"^3\.4\b", "预防性维护-设备保养"),
        (r"^3\.5\b", "故障维修与异常工单"),
        (r"^3\.6\b", "工装管理（详细设计来源）"),
        (r"^3\.7\b", "备件管理"),
    ]:
        if re.match(pattern, text):
            return module
    return None


def esc(text):
    return (text or "").replace("|", "\\|").replace("\r", " ").replace("\n", " / ")


def file_name(index, module):
    return f"{index:02d}-{module}.md".replace("（", "(").replace("）", ")")


def business_name(table_name):
    name = table_name
    replacements = [
        ("oee_", "OEE "),
        ("eq_", "设备 "),
        ("mould_", "工装/模具 "),
        ("spare_parts_", "备件 "),
    ]
    for prefix, label in replacements:
        if name.startswith(prefix):
            name = label + name[len(prefix) :]
            break
    return name.replace("_", " ")


def key_fields(rows):
    keys = []
    preferred = {
        "id",
        "equipment_id",
        "order_id",
        "repair_order_id",
        "task_code",
        "plan_code",
        "maintenance_code",
        "status",
        "state",
        "type",
        "code",
    }
    for row in rows[1:]:
        field = row[1] if len(row) > 1 else ""
        desc = row[7] if len(row) > 7 else ""
        if (
            field in preferred
            or "编号" in desc
            or "状态" in desc
            or "ID" in desc
            or "id" in desc
        ):
            keys.append(f"{field}（{desc}）" if desc else field)
    return "；".join(keys[:8]) if keys else "以字段清单为准"


def markdown_table(rows):
    headers = ["编号", "字段名", "数据类型", "长度", "小数位", "允许空值", "默认值", "说明"]
    lines = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join(["---"] * len(headers)) + "|",
    ]
    for row in rows[1:]:
        normalized = (row + [""] * 8)[:8]
        lines.append("| " + " | ".join(esc(cell) for cell in normalized) + " |")
    return "\n".join(lines)


def collect_tables():
    document = docx.Document(SOURCE)
    blocks = list(iter_blocks(document))
    infos = []
    current_module = "未分组"

    for block_index, (kind, table_index, _) in enumerate(blocks):
        if kind == "p":
            current_module = module_from_heading(blocks[block_index][2]) or current_module
            continue
        if kind != "t" or table_index == 0:
            continue

        previous_paragraphs = [
            block[2]
            for block in blocks[max(0, block_index - 12) : block_index]
            if block[0] == "p" and block[2]
        ]
        table_names = []
        for text in previous_paragraphs:
            match = re.search(r"表[:：]\s*([A-Za-z0-9_]+)", text)
            if match:
                table_names.append(match.group(1))
        name = table_names[-1] if table_names else f"table_{table_index}"

        rows = []
        for row in document.tables[table_index].rows:
            rows.append([cell.text.replace("\n", " / ").strip() for cell in row.cells])
        if not rows or rows[0][:2] != ["编号", "名称"]:
            continue

        infos.append(
            {
                "name": name,
                "module": current_module,
                "rows": rows,
            }
        )
    return infos


def write_index(infos):
    lines = [
        "# 表结构设计索引",
        "",
        "> 来源：`01-详细设计文档1024.docx` 的“数据库设计”章节。本文档集用于帮助产品理解当前系统实现中的业务对象、字段承载和模块边界；不直接等同于新标准产品的目标模型。",
        "",
        "## 使用说明",
        "",
        "- 每个模块文件包含“表清单”和“字段明细”。",
        "- 字段明细保留原详细设计中的字段名、类型、长度、是否允许空值、默认值和说明。",
        "- “产品可读对象/关键字段线索”是基于表名与字段说明的产品化解读，后续项目落地前仍需与研发和现场数据流核对。",
        "- 工装管理来自详细设计文档，但当前需求分析目录未设置同名核心模块，因此作为实现补充单独列出。",
        "",
        "## 模块分布",
        "",
        "| 模块 | 表数量 | 文档 |",
        "|---|---:|---|",
    ]
    for index, module in enumerate(MODULE_ORDER, 1):
        module_tables = [info for info in infos if info["module"] == module]
        name = file_name(index, module)
        lines.append(f"| {module} | {len(module_tables)} | [{name}](./{name}) |")

    lines += [
        "",
        "## 全量表清单",
        "",
        "| 模块 | 表名 | 字段数 | 关键字段线索 |",
        "|---|---|---:|---|",
    ]
    for module in MODULE_ORDER:
        for info in [item for item in infos if item["module"] == module]:
            lines.append(
                f"| {module} | `{info['name']}` | {len(info['rows']) - 1} | {esc(key_fields(info['rows']))} |"
            )
    (OUT_DIR / "00-表结构索引.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_module_files(infos):
    for index, module in enumerate(MODULE_ORDER, 1):
        module_tables = [info for info in infos if info["module"] == module]
        lines = [
            f"# {module}表结构设计",
            "",
            "> 来源：`01-详细设计文档1024.docx`。本文件按业务模块沉淀当前实现表结构，便于产品理解系统现状和后续项目落地差异评估。",
            "",
            "## 表清单",
            "",
            "| 表名 | 产品可读对象 | 字段数 | 关键字段/关联线索 |",
            "|---|---|---:|---|",
        ]
        for info in module_tables:
            lines.append(
                f"| `{info['name']}` | {business_name(info['name'])} | {len(info['rows']) - 1} | {esc(key_fields(info['rows']))} |"
            )
        lines += ["", "## 字段明细", ""]

        for info in module_tables:
            lines += [
                f"### {info['name']}",
                "",
                f"**业务用途（初步解读）**：用于承载“{business_name(info['name'])}”相关数据；具体业务含义以字段说明和详细设计流程为准。",
                "",
                f"**关键字段线索**：{esc(key_fields(info['rows']))}",
                "",
                markdown_table(info["rows"]),
                "",
            ]

        (OUT_DIR / file_name(index, module)).write_text("\n".join(lines), encoding="utf-8")


def main():
    OUT_DIR.mkdir(exist_ok=True)
    infos = collect_tables()
    write_index(infos)
    write_module_files(infos)
    print(f"generated {len(infos)} tables into {OUT_DIR.resolve()}")


if __name__ == "__main__":
    main()
